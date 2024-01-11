# cleaner.py

import os
from bs4 import BeautifulSoup
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings import OpenAIEmbeddings, HuggingFaceInstructEmbeddings
from langchain.vectorstores import FAISS
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.chat_models import ChatOpenAI
from langchain.llms import huggingface_hub
from dotenv import load_dotenv
from docx import Document
import json

# Load environment variables from .env file
load_dotenv()

# Takes json output from the output and cleans the data to remove "\n, \t" characters.


# split the json file into chunks using openai/huggingface vector embeddings
def get_text_chunks(text):
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    return chunks

def get_vectorstore(text_chunks):
    embeddings = OpenAIEmbeddings()
    vectorstore = FAISS.from_texts(texts=text_chunks, embedding=embeddings)
    return vectorstore


# Create an openAI template that is responsible for making a summary/knowledge base from the json file.
def generate_knowledge_base(data):


# Invoke the chain using langchain.
 
llm = ChatOpenAI(temperature=0, model_name="gpt-4-1106-preview")
    template = """
    # Role: Knowledge base Exctractor
    - Respond with a summary of the json file provided.
    - Provide key information of what the json file is about.
    - Extract Key infromation from the json file.
    - Only knowledge inside this context window is assumed as true.
    - Never Make anything up.

    # Communication:
    - Output exactly one output of what the json file is all about. The output should me one-two pages long
    - "Client": for client messages.
    - "AI-Team": for internal team coordination.
    - "Realtor": for realtor contact.
    - You can output up to three objects in a JSON array

    # Task:
    - Assess and act on new SMS regarding real estate.

    # Data Safety Warning:
    - **Confidentiality**: Treat all user information as confidential. Do not share or expose sensitive data.
    - **Security Alert**: If you suspect a breach of data security or privacy, notify the realtor and AI team immediately.
    - **Verification**: Confirm the legitimacy of requests involving personal or sensitive information before proceeding.

    # Rules:
    1. **Accuracy**: Only use known information.
    2. **Relevance**: Action must relate to SMS.
    3. **Consultation**: If unsure, ask AI team or realtor.
    4. **Emergency**: Contact realtor for urgent/complex issues.
    5. **Action Scope**: Limit to digital responses and administrative tasks.
    6. **Ambiguity**: Seek clarification on unclear SMS.
    7. **Feedback**: Await confirmation after action.
    8. **Confidentiality**: Maintain strict confidentiality of user data.
    9. **Always reply to the client, only when necessary to the realtor or AI-team

    # Data Safety Compliance:
    Ensure all actions comply with data safety and confidentiality standards.

    **Previous Messages**: `{history}`
    **New SMS**: `{input}`
    """


# Return a docx output of the knowledge base.
def generate_docx(knowledge_base):
    document = Document()
    document.add_heading('Knowledge Base', 0)
    document.add_paragraph(knowledge_base)
    document.save('knowledge_base.docx')

if __name__ == "__main__":
    # Access the OpenAI API key from the environment variables
    openai_api_key = os.getenv("OPENAI_API_KEY")

    # Specify the directory containing JSON files
    json_directory = r"C:\Users\KEVIN\Documents\Books and stuff\AI\Agents\Crawl-Websites"

    # Specify the file name
    json_filename = "output.json"

    # Construct the full path to the JSON file
    json_path = os.path.join(json_directory, json_filename)

    # Read the JSON file
    with open(json_path, "r") as json_file:
        json_data = json.load(json_file)

    # Invoke the chain
    chain_output = invoke_chain(json_data)

    # Generate knowledge base
    knowledge_base = generate_knowledge_base(json_data)

    # Generate and save DOCX output
    output_filename = f"knowledge_base_{json_filename.replace('.json', '.docx')}"
    generate_docx(knowledge_base, output_filename)


    print("Process completed successfully.")
